#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('03_documentation/Laporan_Praktikum_Data_Science_Customer_Churn_Prediction_NEW.docx')

print("Adding BAB I - PENDAHULUAN (with real data context)...")

def add_black_paragraph(doc, text, indent=True, justify=True):
    """Add paragraph with black font"""
    p = doc.add_paragraph(text)
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_black_heading(doc, text, level):
    """Add heading with black font"""
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

# Page break
doc.add_page_break()

# BAB I
h1 = add_black_heading(doc, 'BAB I', 1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2 = add_black_heading(doc, 'PENDAHULUAN', 1)
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1.1 Business Understanding (EXPLICIT)
add_black_heading(doc, '1.1 Business Understanding (Pemahaman Bisnis)', 2)

add_black_paragraph(doc, 
    'Business Understanding atau pemahaman bisnis adalah tahap pertama dalam siklus '
    'Data Science yang bertujuan memahami konteks bisnis, permasalahan yang dihadapi, '
    'dan tujuan yang ingin dicapai dari perspektif bisnis. Tahap ini penting karena '
    'menentukan arah analisis dan memastikan solusi teknis yang dihasilkan selaras '
    'dengan kebutuhan bisnis.'
)

add_black_heading(doc, '1.1.1 Latar Belakang Masalah', 3)

add_black_paragraph(doc,
    'Industri telekomunikasi merupakan sektor yang sangat kompetitif dengan tingkat '
    'persaingan tinggi. Customer churn (istilah bahasa Inggris yang berarti perpindahan '
    'atau kehilangan pelanggan) menjadi tantangan kritis yang dihadapi perusahaan telekomunikasi. '
    'Customer churn didefinisikan sebagai kondisi ketika pelanggan menghentikan layanan '
    'dari satu provider dan beralih ke kompetitor lain.'
)

add_black_paragraph(doc,
    'Berdasarkan penelitian Ahmad et al. (2019) dalam jurnal "Customer churn prediction '
    'in telecom using machine learning in big data platform" yang diterbitkan di Journal '
    'of Big Data (Springer), ditemukan bahwa biaya untuk mendapatkan pelanggan baru '
    '(Customer Acquisition Cost atau CAC) mencapai enam kali lipat lebih mahal dibandingkan '
    'biaya untuk mempertahankan pelanggan yang sudah ada (Customer Retention Cost). '
    'Hal ini membuat strategi retensi pelanggan menjadi prioritas utama dari segi finansial.'
)

add_black_paragraph(doc,
    'Dalam konteks bisnis telekomunikasi, kehilangan pelanggan berdampak langsung pada '
    'penurunan revenue atau pendapatan perusahaan. Setiap pelanggan yang churn '
    'merepresentasikan kehilangan revenue bulanan berulang (recurring revenue) yang '
    'signifikan. Oleh karena itu, kemampuan untuk memprediksi pelanggan mana yang '
    'berpotensi melakukan churn menjadi sangat valuable atau bernilai tinggi bagi '
    'perusahaan.'
)

add_black_heading(doc, '1.1.2 Objektif Bisnis', 3)

add_black_paragraph(doc, 'Berdasarkan analisis kebutuhan bisnis, praktikum ini memiliki objektif sebagai berikut:')

objectives = [
    'Mengembangkan model prediksi (prediction model) yang dapat mengidentifikasi pelanggan '
    'berisiko tinggi melakukan churn dengan tingkat akurasi yang baik (target minimum 75%)',
    
    'Mengidentifikasi faktor-faktor atau variabel (features) yang paling berpengaruh terhadap '
    'keputusan pelanggan untuk melakukan churn, sehingga perusahaan dapat mengambil tindakan preventif',
    
    'Menyediakan early warning system (sistem peringatan dini) yang memungkinkan tim marketing '
    'dan customer service melakukan intervensi proaktif kepada pelanggan berisiko tinggi',
    
    'Mengoptimalkan alokasi budget marketing dengan fokus pada pelanggan yang benar-benar '
    'memerlukan retention program, sehingga efisiensi biaya dapat tercapai'
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(f'{i}. {obj}')
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

# 1.2 Data Understanding (EXPLICIT)
add_black_heading(doc, '1.2 Data Understanding (Pemahaman Data)', 2)

add_black_paragraph(doc,
    'Data Understanding adalah tahap kedua dalam siklus Data Science di mana kita melakukan '
    'eksplorasi terhadap data yang tersedia, memahami karakteristiknya, mengidentifikasi '
    'kualitas data, dan menemukan insight awal. Tahap ini krusial karena menentukan '
    'kualitas model yang akan dibangun - prinsip "garbage in, garbage out" berlaku di sini.'
)

add_black_heading(doc, '1.2.1 Sumber Data', 3)

add_black_paragraph(doc,
    'Praktikum ini menggunakan dataset real (data nyata) dari Kaggle yaitu IBM Telco '
    'Customer Churn Dataset. Kaggle adalah platform komunitas data science terbesar '
    'di dunia yang menyediakan dataset untuk pembelajaran dan kompetisi. Dataset ini '
    'dipilih karena:'
)

reasons = [
    'Merupakan data real-world dari industri telekomunikasi yang actual',
    'Memiliki ukuran yang cukup representatif (7,043 customer records)',
    'Mencakup berbagai dimensi informasi pelanggan (demografis, layanan, billing)',
    'Banyak digunakan dalam penelitian akademis dan industri sebagai benchmark',
    'Tersedia secara publik dan dapat diverifikasi'
]

for reason in reasons:
    p = doc.add_paragraph(reason, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

print(f"BAB I added. Total paragraphs: {len(doc.paragraphs)}")

# Save
doc.save('03_documentation/Laporan_Praktikum_Data_Science_Customer_Churn_Prediction_NEW.docx')
print("Progress saved")

