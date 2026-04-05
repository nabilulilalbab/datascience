#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('03_documentation/Laporan_Praktikum_Data_Science_Customer_Churn_Prediction_NEW.docx')

print("Adding BAB II - LANDASAN TEORI (super detail explanations)...")

def add_black_paragraph(doc, text, indent=True):
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_black_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

doc.add_page_break()

# BAB II
h1 = add_black_heading(doc, 'BAB II', 1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2 = add_black_heading(doc, 'LANDASAN TEORI', 1)
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 2.1 Data Science
add_black_heading(doc, '2.1 Data Science', 2)

add_black_paragraph(doc,
    'Data Science adalah disiplin ilmu yang menggabungkan metode statistika, matematika, '
    'dan ilmu komputer untuk mengekstrak pengetahuan dan insight dari data. Secara sederhana, '
    'Data Science adalah proses mengubah data mentah menjadi informasi yang berguna untuk '
    'pengambilan keputusan. Menurut Ajah dan Nweke (2019) dalam jurnal "Big data and business '
    'analytics: Trends, platforms, success factors and applications" yang diterbitkan di '
    'Big Data and Cognitive Computing (MDPI), Data Science memiliki tiga komponen utama:'
)

components = [
    'Descriptive Analytics: Analisis deskriptif yang menjawab pertanyaan "apa yang terjadi?" '
    'melalui statistika deskriptif dan visualisasi data',
    
    'Predictive Analytics: Analisis prediktif yang menjawab pertanyaan "apa yang akan terjadi?" '
    'menggunakan machine learning dan statistical modeling',
    
    'Prescriptive Analytics: Analisis preskriptif yang menjawab pertanyaan "apa yang harus dilakukan?" '
    'dengan memberikan rekomendasi tindakan optimal'
]

for comp in components:
    p = doc.add_paragraph(comp, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

# 2.2 Machine Learning
add_black_heading(doc, '2.2 Machine Learning', 2)

add_black_paragraph(doc,
    'Machine Learning (pembelajaran mesin) adalah cabang dari Artificial Intelligence (kecerdasan '
    'buatan) yang memungkinkan komputer untuk belajar dari data tanpa diprogram secara eksplisit '
    'untuk setiap kasus. Berbeda dengan pemrograman tradisional di mana programmer menulis aturan '
    'secara manual, dalam machine learning komputer "belajar" pola dari data dan membuat prediksi '
    'atau keputusan berdasarkan pola tersebut.'
)

add_black_paragraph(doc,
    'Machine Learning dibagi menjadi tiga kategori utama berdasarkan cara pembelajaran:'
)

ml_types = [
    'Supervised Learning (pembelajaran terawasi): Algoritma belajar dari data yang sudah memiliki '
    'label atau jawaban yang benar. Contohnya adalah prediksi churn di mana kita sudah tahu pelanggan '
    'mana yang churn dan tidak churn di masa lalu. Model belajar dari contoh-contoh ini.',
    
    'Unsupervised Learning (pembelajaran tak terawasi): Algoritma mencari pola dalam data tanpa label. '
    'Contohnya adalah clustering atau pengelompokan pelanggan berdasarkan kesamaan karakteristik.',
    
    'Reinforcement Learning (pembelajaran penguatan): Algoritma belajar melalui trial and error '
    'dengan sistem reward dan penalty. Sering digunakan untuk game AI dan robotika.'
]

for ml_type in ml_types:
    p = doc.add_paragraph(ml_type, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

add_black_paragraph(doc,
    'Praktikum ini menggunakan Supervised Learning karena kita memiliki data historis pelanggan '
    'dengan label churn (Yes/No) yang sudah diketahui.'
)

# 2.3 Random Forest  
add_black_heading(doc, '2.3 Random Forest Classifier', 2)

add_black_paragraph(doc,
    'Random Forest adalah algoritma machine learning yang termasuk dalam kategori ensemble learning. '
    'Ensemble learning adalah teknik yang menggabungkan prediksi dari beberapa model untuk menghasilkan '
    'prediksi yang lebih akurat dan robust (tahan terhadap error) dibandingkan model tunggal.'
)

add_black_paragraph(doc,
    'Cara kerja Random Forest dapat dijelaskan dengan analogi voting demokratis: bayangkan kita memiliki '
    '100 orang ahli yang masing-masing memberikan pendapat tentang apakah seorang pelanggan akan churn. '
    'Keputusan akhir diambil berdasarkan suara terbanyak (majority voting). Dalam Random Forest, '
    '"ahli" tersebut adalah Decision Tree (pohon keputusan).'
)

add_black_paragraph(doc,
    'Komponen Random Forest:'
)

rf_components = [
    'Decision Tree: Model prediksi berbentuk pohon dengan cabang-cabang yang merepresentasikan keputusan '
    'berdasarkan nilai fitur tertentu. Contoh: "jika tenure < 12 bulan DAN MonthlyCharges > 70, maka churn"',
    
    'Bootstrap Sampling: Teknik mengambil sampel data secara acak dengan replacement (data yang sama '
    'bisa terpilih lebih dari sekali) untuk melatih setiap tree',
    
    'Feature Randomness: Setiap tree hanya menggunakan subset acak dari fitur yang tersedia, bukan semua fitur. '
    'Ini membuat setiap tree berbeda dan mengurangi correlation antar tree',
    
    'Voting Mechanism: Untuk klasifikasi, prediksi akhir adalah kelas dengan suara terbanyak dari semua tree'
]

for i, comp in enumerate(rf_components, 1):
    p = doc.add_paragraph(f'{i}. {comp}')
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

add_black_paragraph(doc,
    'Menurut Imani et al. (2025) dalam systematic review mereka yang berjudul "Customer churn prediction: '
    'A systematic review of recent advances" di jurnal Machine Learning and Knowledge Extraction (MDPI), '
    'Random Forest menunjukkan performa yang sangat baik untuk prediksi churn karena beberapa keunggulan: '
    '(1) dapat menangani imbalanced data (ketidakseimbangan kelas) dengan baik, (2) resistant terhadap '
    'overfitting (model terlalu fokus pada data training sehingga gagal memprediksi data baru), dan '
    '(3) dapat mengukur feature importance (tingkat kepentingan setiap variabel).'
)

print(f"BAB II partial added. Total paragraphs: {len(doc.paragraphs)}")
doc.save('03_documentation/Laporan_Praktikum_Data_Science_Customer_Churn_Prediction_NEW.docx')
print("Progress saved (BAB II part 1/2)")

