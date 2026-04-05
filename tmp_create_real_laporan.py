#!/usr/bin/env python3
"""
Create completely new Laporan Praktikum with REAL data (7,043 samples)
Super detailed explanations for every technical term
Explicit Business Understanding & Data Understanding
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*79)
print("CREATING NEW LAPORAN WITH REAL DATA")
print("="*79)

# Create new document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

print("\nStep 1: Creating cover page...")

# Cover Page
title = doc.add_heading('LAPORAN PRAKTIKUM DATA SCIENCE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_paragraph('Prediksi Customer Churn pada Industri Telekomunikasi')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 0)

subtitle2 = doc.add_paragraph('Menggunakan Random Forest Classifier')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(12)
subtitle2.runs[0].font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ['Mata Kuliah', 'Data Science'],
    ['Topik', 'Customer Churn Prediction'],
    ['Dataset', 'IBM Telco Customer Churn (Kaggle)'],
    ['Ukuran Dataset', '7,043 customer records'],
    ['Algoritma', 'Random Forest Classifier'],
    ['Tanggal', '5 April 2026']
]
for i, (key, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = key
    info_table.rows[i].cells[1].text = value
    # Set font color to black
    for cell in info_table.rows[i].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

print("Cover page created")

# Save progress
doc.save('03_documentation/Laporan_Praktikum_Data_Science_Customer_Churn_Prediction_NEW.docx')
print("\nProgress saved: Laporan_Praktikum_NEW.docx")
print("Total paragraphs so far:", len(doc.paragraphs))

